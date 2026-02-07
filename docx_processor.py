"""
DOCX Processor - OPTIMIZED
--------------------------
Fast processing with batch translation.
"""

import logging
from typing import List, Tuple
from docx import Document
from docx.table import Table

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXProcessor:
    def __init__(self, translator, status_callback=None):
        self.translator = translator
        self.status_callback = status_callback
        self.stats = {
            'paragraphs_processed': 0,
            'tables_processed': 0,
            'headers_processed': 0,
            'footers_processed': 0,
            'text_runs_translated': 0,
            'errors': []
        }
    
    def _update_status(self, message):
        if self.status_callback:
            self.status_callback(message)
        logger.info(message)
    
    def process_file(self, input_path: str, output_path: str) -> dict:
        try:
            self._update_status("Loading document...")
            doc = Document(input_path)
            
            # Collect all text locations
            text_locations: List[Tuple] = []
            
            # Main body
            self._update_status("Collecting text from document...")
            for para in doc.paragraphs:
                self._collect_paragraph(para, text_locations)
                self.stats['paragraphs_processed'] += 1
            
            # Tables
            for table in doc.tables:
                self._collect_table(table, text_locations)
                self.stats['tables_processed'] += 1
            
            # Headers and footers
            for section in doc.sections:
                try:
                    if section.header:
                        for para in section.header.paragraphs:
                            self._collect_paragraph(para, text_locations)
                        self.stats['headers_processed'] += 1
                    if section.footer:
                        for para in section.footer.paragraphs:
                            self._collect_paragraph(para, text_locations)
                        self.stats['footers_processed'] += 1
                except:
                    pass
            
            if text_locations:
                self._update_status(f"Batch translating {len(text_locations)} text blocks...")
                
                # Batch translate
                all_texts = [loc[2] for loc in text_locations]
                translated = self.translator.translate_batch(all_texts)
                
                # Apply translations
                for (para, runs, orig_combined, orig_texts), trans in zip(text_locations, translated):
                    if trans and trans != orig_combined:
                        self._redistribute_text_to_runs(runs, orig_texts, trans)
                        self.stats['text_runs_translated'] += len(runs)
            
            self._update_status("Saving translated document...")
            doc.save(output_path)
            
            self.stats['translator_stats'] = self.translator.get_stats()
            return self.stats
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            self.stats['errors'].append(str(e))
            raise
    
    def _collect_paragraph(self, para, text_locations: List) -> None:
        """Collect text from a paragraph."""
        runs = list(para.runs)
        if not runs:
            return
        
        orig_texts = [r.text for r in runs]
        combined = ''.join(orig_texts)
        
        if combined.strip() and len(combined.strip()) >= 2:
            text_locations.append((para, runs, combined, orig_texts))
    
    def _collect_table(self, table: Table, text_locations: List) -> None:
        """Collect text from table cells."""
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._collect_paragraph(para, text_locations)
        except Exception as e:
            self.stats['errors'].append(f"Table error: {e}")
    
    def _redistribute_text_to_runs(self, runs, original_texts, translated_text) -> None:
        """Redistribute translated text proportionally."""
        total_len = sum(len(t) for t in original_texts)
        
        if total_len == 0:
            if runs:
                runs[0].text = translated_text
            return
        
        trans_len = len(translated_text)
        pos = 0
        
        for i, (run, orig) in enumerate(zip(runs, original_texts)):
            if i == len(runs) - 1:
                run.text = translated_text[pos:]
            else:
                prop = len(orig) / total_len
                chars = int(trans_len * prop)
                end = pos + chars
                
                if end < trans_len:
                    space = translated_text.rfind(' ', pos, end + 10)
                    if space > pos:
                        end = space + 1
                
                run.text = translated_text[pos:end]
                pos = end
